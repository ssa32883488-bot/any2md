#!/usr/bin/env python3
"""Generate any2md stress-test dataset: multi-column / multi-layout / multi-level Word + PDF."""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
import textwrap
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "testset"
OFFICE = OUT / "office"
DIGITAL = OUT / "digital"
SCAN = OUT / "scan"
EXPECTED = OUT / "expected"


def _set_columns(section, count: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "720")


def _add_heading_para(doc, text: str, level: int = 1) -> None:
    try:
        doc.add_heading(text, level=min(level, 9))
    except Exception:
        doc.add_paragraph(text)


def _add_cn_section(doc, title: str, items: list[str], subsections: list[tuple[str, list[str]]] | None = None) -> None:
    p = doc.add_paragraph(title)
    try:
        p.style = "Heading 1"
    except Exception:
        pass
    for item in items:
        doc.add_paragraph(item)
    if subsections:
        for sub_title, sub_items in subsections:
            sp = doc.add_paragraph(sub_title)
            try:
                sp.style = "Heading 2"
            except Exception:
                pass
            for si in sub_items:
                doc.add_paragraph(si)


def _add_data_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val


def build_01_multilevel(doc) -> None:
    doc.add_heading("压力测试 01：多层级标题与列表", 0)
    doc.add_paragraph("用途：验证 Word 标题样式、中文序号、嵌套列表的结构化提取与切分。")
    _add_cn_section(
        doc,
        "一、文档说明",
        [
            "1. 本文件包含三级标题与混合编号列表。",
            "2. 用于测试 CPU 快路径 docx → Markdown。",
            "3. 转换后应保留 ## / ### 层级。",
        ],
    )
    _add_cn_section(
        doc,
        "二、技术要点",
        [],
        [
            (
                "1. 阈值分割原理",
                [
                    "阈值分割通过灰度阈值 T 将前景与背景分离。",
                    "Otsu 算法通过最大化类间方差选取最佳阈值。",
                ],
            ),
            (
                "2. 边缘检测原理",
                [
                    "Sobel / Prewitt 算子利用一阶导数检测边缘。",
                    "Canny 算子包含高斯平滑、非极大值抑制与双阈值连接。",
                ],
            ),
        ],
    )
    _add_cn_section(doc, "三、结论", ["层级结构应完整出现在 Markdown 与 chunks 中。"])


def build_02_two_column(doc) -> None:
    doc.add_heading("压力测试 02：双栏版面", 0)
    doc.add_paragraph("第一页为单栏封面说明。")
    doc.add_page_break()
    sec = doc.add_section()
    _set_columns(sec, 2)
    doc.add_heading("双栏正文开始", level=1)
    for i in range(1, 9):
        doc.add_paragraph(
            f"【栏块 {i}】双栏布局常见于论文与期刊。"
            "本段用于测试分栏段落在转换时是否被正确展开、顺序是否保持。"
            "若顺序错乱或合并为一行，则说明解析器需增强分栏支持。"
        )
    doc.add_page_break()
    sec2 = doc.add_section()
    _set_columns(sec2, 1)
    doc.add_heading("恢复单栏", level=1)
    doc.add_paragraph("分栏结束后应回到单栏流式排版。")


def build_03_three_column(doc) -> None:
    doc.add_heading("压力测试 03：三栏紧凑布局", 0)
    sec = doc.sections[0]
    _set_columns(sec, 3)
    for c in range(1, 13):
        doc.add_paragraph(f"三栏段落 {c}：短文本块，模拟词典/索引式排版。")


def build_04_layout_table(doc) -> None:
    """WPS/Word 模板常见：正文放在 1×1 布局表格内。"""
    doc.add_heading("压力测试 04：布局表格包裹正文", 0)
    doc.add_paragraph("《综合实验报告》")
    doc.add_paragraph("专业班级：测试班 2026")
    doc.add_paragraph("姓名：any2md 测试")
    _add_data_table(
        doc,
        ["序号", "评价项", "等级"],
        [["1", "结构完整", "优"], ["2", "格式规范", "良"]],
    )
    body = doc.add_table(rows=1, cols=1)
    cell = body.cell(0, 0)
    cell.paragraphs[0].text = "一、实验目的"
    cell.add_paragraph("1. 验证布局表格内段落能否识别为标题。")
    cell.add_paragraph("2. 验证列表项 1、2、3 不被误判为标题。")
    cell.add_paragraph("二、实验原理")
    p = cell.add_paragraph("1. 采样定理")
    if p.runs:
        p.runs[0].bold = True
    cell.add_paragraph("奈奎斯特采样定理是数字信号处理的基础。")
    cell.add_paragraph("三、实验步骤")
    cell.add_paragraph("连接设备，采集数据，绘制频谱图。")
    cell.add_paragraph("四、结论")
    cell.add_paragraph("布局表格文档应输出带 ## 层级的 Markdown。")


def build_05_mixed_stress(doc) -> None:
    doc.add_heading("压力测试 05：混合大文档", 0)
    doc.add_paragraph(f"生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_paragraph("本文件页数较多，用于批量转换与切分压力测试。")

    for ch in range(1, 6):
        doc.add_page_break()
        _add_cn_section(
            doc,
            f"{_cn_num(ch)}、章节 {ch}：混合内容",
            [f"章节 {ch} 概述段落。" * 3],
            [
                (
                    f"{sub}. 小节 {ch}.{sub}",
                    [
                        f"小节正文 {ch}.{sub}：" + "测试句子。" * 20,
                        f"补充说明 {ch}.{sub}：" + "更多内容。" * 15,
                    ],
                )
                for sub in range(1, 4)
            ],
        )
        _add_data_table(
            doc,
            ["参数", "数值", "单位"],
            [[f"P{i}", str(i * 3.14)[:6], "mm"] for i in range(1, 6)],
        )
        if ch % 2 == 0:
            sec = doc.add_section()
            _set_columns(sec, 2)
            for j in range(1, 5):
                doc.add_paragraph(f"章节 {ch} 双栏段 {j}：" + "分栏文本。" * 12)
            doc.add_section()

    doc.add_page_break()
    doc.add_heading("附录 A：参考文献", level=1)
    for i in range(1, 11):
        doc.add_paragraph(f"[{i}] 参考文献条目 {i}，作者 et al., 202{i % 5}.")


def build_06_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "参数表"
    ws1.append(["ID", "名称", "阈值", "备注"])
    for i in range(1, 21):
        ws1.append([i, f"传感器-{i}", i * 0.05, "测试数据"])
    ws2 = wb.create_sheet("统计")
    ws2.append(["类别", "数量"])
    ws2.append(["A类", 12])
    ws2.append(["B类", 8])
    wb.save(path)


def _cn_num(n: int) -> str:
    chars = "零一二三四五六七八九十"
    if n <= 10:
        return chars[n]
    if n < 20:
        return "十" + chars[n - 10]
    return str(n)


def _extract_docx_blocks(path: Path) -> list[tuple[str, str]]:
    """Extract (kind, text) from docx for PDF rendering. kind: h1/h2/p/t."""
    from docx import Document

    sys.path.insert(0, str(ROOT / "engine"))
    from docx_structured import iter_block_items, _para_to_md, _blocks_from_table  # noqa: E402

    doc = Document(path)
    blocks: list[tuple[str, str]] = []

    def flush_para(para) -> None:
        md = _para_to_md(para)
        if not md:
            return
        if md.startswith("## "):
            blocks.append(("h1", md[3:]))
        elif md.startswith("### "):
            blocks.append(("h2", md[4:]))
        elif md.startswith("# "):
            blocks.append(("h0", md[2:]))
        else:
            blocks.append(("p", md))

    for item in iter_block_items(doc):
        if hasattr(item, "text"):
            flush_para(item)
        else:
            for part in _blocks_from_table(item):
                if part.startswith("|"):
                    blocks.append(("table", part))
                elif part.startswith("#"):
                    if part.startswith("### "):
                        blocks.append(("h2", part[4:]))
                    elif part.startswith("## "):
                        blocks.append(("h1", part[3:]))
                    else:
                        blocks.append(("h0", part.lstrip("# ").strip()))
                else:
                    blocks.append(("p", part))
    return blocks


def docx_to_digital_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Render a digital PDF mirror (no MS Word required)."""
    import fitz

    blocks = _extract_docx_blocks(docx_path)
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    y = 50
    margin = 50
    width = 495

    def new_page_if(need: float = 40) -> None:
        nonlocal page, y
        if y + need > 800:
            page = doc.new_page(width=595, height=842)
            y = 50

    for kind, text in blocks:
        if kind == "table":
            new_page_if(60)
            for line in text.splitlines()[:12]:
                page.insert_text((margin, y), line[:100], fontsize=8, fontname="china-s")
                y += 11
            y += 8
            continue

        if kind == "h0":
            new_page_if(36)
            page.insert_text((margin, y), text[:80], fontsize=18, fontname="china-s")
            y += 28
        elif kind == "h1":
            new_page_if(30)
            page.insert_text((margin, y), text[:80], fontsize=14, fontname="china-s")
            y += 22
        elif kind == "h2":
            new_page_if(26)
            page.insert_text((margin + 8, y), text[:80], fontsize=12, fontname="china-s")
            y += 18
        else:
            wrapped = textwrap.wrap(text, width=42) or [text]
            for line in wrapped:
                new_page_if(14)
                page.insert_text((margin, y), line, fontsize=10, fontname="china-s")
                y += 13
            y += 6

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(pdf_path)
    doc.close()


def try_office_export(docx_path: Path, pdf_path: Path) -> str | None:
    """Try Word COM export; return method name or None."""
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return None

    for prog_id in ("Word.Application", "Kwps.Application", "Wps.Application"):
        try:
            app = win32com.client.Dispatch(prog_id)
            app.Visible = False
            doc = app.Documents.Open(str(docx_path.resolve()))
            doc.ExportAsFixedFormat(str(pdf_path.resolve()), 17)  # wdExportFormatPDF
            doc.Close(False)
            app.Quit()
            return prog_id
        except Exception:
            try:
                app.Quit()
            except Exception:
                pass
    return None


def build_manifest(files: list[dict]) -> None:
    manifest = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "purpose": "any2md 压力测试集：多栏、多版面、多层级",
        "files": files,
        "scan_instructions": (
            "打开 digital/*.pdf 或 office/*.pdf，使用 Win+Shift+S 或扫描仪逐页截图，"
            "合并为 scan/*.pdf，用于 OCR 路径测试。"
        ),
    }
    (OUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    from docx import Document

    for d in (OFFICE, DIGITAL, SCAN, EXPECTED):
        d.mkdir(parents=True, exist_ok=True)

    specs: list[tuple[str, callable]] = [
        ("01_multilevel_sections", build_01_multilevel),
        ("02_two_column", build_02_two_column),
        ("03_three_column", build_03_three_column),
        ("04_layout_table_body", build_04_layout_table),
        ("05_mixed_stress", build_05_mixed_stress),
    ]

    manifest_files: list[dict] = []

    for stem, builder in specs:
        docx_path = OFFICE / f"{stem}.docx"
        pdf_office = OFFICE / f"{stem}.pdf"
        pdf_digital = DIGITAL / f"{stem}.pdf"

        doc = Document()
        builder(doc)
        doc.save(docx_path)

        export_via = try_office_export(docx_path, pdf_office)
        if not export_via or not pdf_office.is_file():
            docx_to_digital_pdf(docx_path, pdf_office)
            export_via = "pymupdf-from-docx"

        docx_to_digital_pdf(docx_path, pdf_digital)

        manifest_files.append(
            {
                "id": stem,
                "docx": str(docx_path.relative_to(ROOT)),
                "pdf_office": str(pdf_office.relative_to(ROOT)),
                "pdf_digital": str(pdf_digital.relative_to(ROOT)),
                "pdf_source": export_via,
                "route_hint": {
                    "docx": "text (CPU fast)",
                    "pdf_office": "text if digital / ocr if scanned",
                    "pdf_digital": "text (CPU fast)",
                },
                "scan_target": str((SCAN / f"{stem}_scan.pdf").relative_to(ROOT)),
            }
        )
        print(f"OK {stem}: docx + pdf ({export_via})")

    xlsx_path = OFFICE / "06_excel_tables.xlsx"
    build_06_xlsx(xlsx_path)
    manifest_files.append(
        {
            "id": "06_excel_tables",
            "xlsx": str(xlsx_path.relative_to(ROOT)),
            "route_hint": "text (CPU fast)",
        }
    )
    print(f"OK 06_excel_tables.xlsx")

    build_manifest(manifest_files)

    (SCAN / "README.txt").write_text(
        textwrap.dedent(
            """
            扫描件 PDF 制作说明（需人工完成，用于 OCR 压力测试）
            ==================================================

            1. 打开 testset/digital/ 或 testset/office/ 下的 PDF
            2. 放大到 100%–125%，确保文字清晰
            3. 逐页截图（Win+Shift+S）或使用扫描仪
            4. 将图片合并为 PDF，保存为本目录下同名文件，例如：
               - 01_multilevel_sections_scan.pdf
               - 05_mixed_stress_scan.pdf
            5. 在 any2md 中选择「自动」或「OCR」模式转换

            建议至少制作 2 份扫描件：
            - 01_multilevel_sections_scan.pdf（层级）
            - 05_mixed_stress_scan.pdf（页数多，压力）
            """
        ).strip(),
        encoding="utf-8",
    )

    (EXPECTED / "checklist.md").write_text(
        textwrap.dedent(
            """
            # 测试检查清单

            | 文件 | 期望 Markdown 特征 | 切分期望 |
            |------|-------------------|----------|
            | 01_multilevel_sections | ## 一、二、三、 | 按章节分块 |
            | 02_two_column | 双栏段落顺序正确 | 段落/章节块 |
            | 03_three_column | 三栏短段落 | 多块 |
            | 04_layout_table_body | 表格+## 标题，非单行巨表 | 按 ## 切分 |
            | 05_mixed_stress | 5 章+附录，含表格 | 多块，batch 测试 |
            | 06_excel_tables | 多 sheet 表格 | 可选切分 |
            | *_scan.pdf | OCR 路径 | 与 digital 对比 |
            """
        ).strip(),
        encoding="utf-8",
    )

    readme = OUT / "README.md"
    readme.write_text(
        f"""# any2md 压力测试集

生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}

## 目录

| 目录 | 内容 |
|------|------|
| `office/` | Word / Excel 源文件 + 配套 PDF |
| `digital/` | 数字 PDF（pymupdf 渲染，适合 CPU 快路径） |
| `scan/` | **请自行截图/扫描生成的 PDF 放这里** |
| `expected/` | 验收检查清单 |
| `manifest.json` | 机器可读清单 |

## 快速压测命令

```powershell
cd f:\\code\\any2md
$py = "F:\\Python\\Python 3.13.0\\python.exe"

# 批量 CPU 快路径
Get-ChildItem testset\\office\\*.docx | ForEach-Object {{
  & $py engine/run_parser.py -i $_.FullName -o testset\\output --route text --chunk-model bge-base-zh-v1.5
}}

# 数字 PDF
Get-ChildItem testset\\digital\\*.pdf | ForEach-Object {{
  & $py engine/run_parser.py -i $_.FullName -o testset\\output --route auto
}}

# 扫描件（放入 testset\\scan\\ 后）
Get-ChildItem testset\\scan\\*_scan.pdf | ForEach-Object {{
  & $py engine/run_parser.py -i $_.FullName -o testset\\output --route ocr
}}
```

## 文件说明

{chr(10).join(f"- **{m['id']}**" for m in manifest_files)}
""",
        encoding="utf-8",
    )

    print(f"\nDone -> {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
